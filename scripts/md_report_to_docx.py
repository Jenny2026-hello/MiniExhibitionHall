from pathlib import Path
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
INPUT = Path(sys.argv[1]) if len(sys.argv) > 1 else ROOT / "Docs" / "项目报告.md"
OUTPUT = Path(sys.argv[2]) if len(sys.argv) > 2 else ROOT / "Docs" / "项目报告.docx"

FONT_CN = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 96, 108)


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def add_inline_markdown(paragraph, text, base_size=11, color=None):
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, base_size, True, color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(base_size)
            if color is not None:
                run.font.color.rgb = color
        else:
            run = paragraph.add_run(part)
            set_run_font(run, base_size, None, color)


def add_heading(doc, text, level):
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=0, after=12, line=1.05)
        run = p.add_run(text)
        set_run_font(run, 20, True, DARK_BLUE)
    elif level == 2:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=16, after=8)
        run = p.add_run(text)
        set_run_font(run, 16, True, BLUE)
    elif level == 3:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=12, after=6)
        run = p.add_run(text)
        set_run_font(run, 13, True, BLUE)
    else:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=8, after=4)
        run = p.add_run(text)
        set_run_font(run, 12, True, DARK_BLUE)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    add_inline_markdown(p, text)


def add_list_item(doc, text, ordered=False):
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    set_paragraph_spacing(p, after=4, line=1.10)
    add_inline_markdown(p, text)


def parse_table(lines, start):
    table_lines = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    rows = []
    for idx, line in enumerate(table_lines):
        cells = [c.strip() for c in line.strip("|").split("|")]
        if idx == 1 and all(set(c.replace(":", "")) <= {"-"} for c in cells):
            continue
        rows.append(cells)
    return rows, i


def add_markdown_table(doc, rows):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            text = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.05)
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                add_inline_markdown(p, text, base_size=10)
                for run in p.runs:
                    run.bold = True
            else:
                add_inline_markdown(p, text, base_size=10)
    widths = [6.3 / col_count] * col_count
    if col_count == 2:
        widths = [1.7, 4.6]
    elif col_count == 3:
        widths = [1.1, 2.6, 2.6]
    set_table_width(table, widths)
    doc.add_paragraph()


def configure_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.49)
    sec.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(11)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("MiniExhibitionHall — 太空探索虚拟展览馆")
    set_run_font(run, 9, False, MUTED)


def build_docx():
    doc = Document()
    configure_document(doc)
    lines = INPUT.read_text(encoding="utf-8").splitlines()

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line or line == "---":
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            add_heading(doc, m.group(2).strip(), len(m.group(1)))
        elif line.startswith("- "):
            add_list_item(doc, line[2:].strip(), ordered=False)
        elif re.match(r"^\d+\.\s+", line):
            add_list_item(doc, re.sub(r"^\d+\.\s+", "", line).strip(), ordered=True)
        elif line.startswith(">"):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=4, after=8, line=1.10)
            add_inline_markdown(p, line.lstrip("> ").strip(), color=MUTED)
        else:
            add_body_paragraph(doc, line)
        i += 1

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    out = build_docx()
    print(out)
